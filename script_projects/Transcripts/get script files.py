import docx
import requests
import re
from youtube_transcript_api import YouTubeTranscriptApi




# Преваратить в функцию парсинг
# Преваратить в функцию форматирование текста

api_key = "AIzaSyBd9-w1yBcoTsE0y30_Olikk4zCPe8SZ-Q"
base_search_url = 'https://www.googleapis.com/youtube/v3/videos?'

video_ids = []
with open("Hubrman video links.txt", "r", encoding='utf-8') as file:
    for line in file.readlines():
        video_ids.append(line.split('watch?v=')[1].replace('\n', ''))

timestamps_doc = docx.Document()

for i, video_id in enumerate(video_ids):
    first_url = base_search_url + 'key={}&part=snippet&id={}'.format(api_key, video_id)
    jsom_url = requests.get(first_url).json()
    video_title = jsom_url['items'][0]['snippet']['title']
    video_title = re.sub('[^\w]', ' ', video_title)
    video_title = video_title.replace('Huberman Lab Podcast', '').strip()
    while '  ' in video_title:
        video_title = video_title.replace('  ', ' ')
    video_description = jsom_url['items'][0]['snippet']['description']
    video_timestamps = ''
    if re.findall('Timestamps', video_description):
        timestamps = video_description.split('Timestamps')[-1]
        splitted_stamps = timestamps.split('\n')
        for line in splitted_stamps:
            if len(line) > 0:
                if line[0].isdigit():
                    video_timestamps += line + '\n'
    else:
        video_timestamps = video_description
    if video_id == 'ufsIA5NARIo':
        transcription_text = ''
    else:
        try:
            transcription = YouTubeTranscriptApi.get_transcript(video_id, languages=['en-US'])
        except:
            transcription = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
        transcription_text = ''
        for part in transcription:
            transcription_text += part['text'] + ' '
        transcription_text = transcription_text.replace('\n', ' ')

    video_tags = jsom_url['items'][0]['snippet']['tags']
    tags_paragraph = ''
    for tag in video_tags:
        tags_paragraph += tag + ', '
    tags_paragraph = tags_paragraph.rstrip(', ')

    p = timestamps_doc.add_paragraph()
    p.paragraph_format.space_after = 0
    p.paragraph_format.line_spacing = 1.5
    p.alignment = 1

    run = p.add_run(video_title)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = docx.shared.Pt(14)

    episode_doc = docx.Document()
    p = episode_doc.add_paragraph()
    p.paragraph_format.space_after = 0
    p.paragraph_format.line_spacing = 1.5
    p.alignment = 1

    run = p.add_run('Tags')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = docx.shared.Pt(14)

    p = episode_doc.add_paragraph()
    p.paragraph_format.space_after = 0
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = docx.shared.Mm(12.5)
    p.alignment = 3

    run = p.add_run(tags_paragraph)
    run.font.name = 'Times New Roman'
    run.font.size = docx.shared.Pt(14)
    episode_doc.add_paragraph()

    p = episode_doc.add_paragraph()
    p.paragraph_format.space_after = 0
    p.paragraph_format.line_spacing = 1.5
    p.alignment = 1

    run = p.add_run('Timestamps')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = docx.shared.Pt(14)

    for line in video_timestamps.split('\n'):
        p = episode_doc.add_paragraph()
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = docx.shared.Mm(12.5)
        p.alignment = 3
        run = p.add_run(line.strip())
        run.font.name = 'Times New Roman'
        run.font.size = docx.shared.Pt(14)

        p = timestamps_doc.add_paragraph()
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = docx.shared.Mm(12.5)
        p.alignment = 3
        run = p.add_run(line.strip())
        run.font.name = 'Times New Roman'
        run.font.size = docx.shared.Pt(14)

    timestamps_doc.add_paragraph()

    p = episode_doc.add_paragraph()
    p.paragraph_format.space_after = 0
    p.paragraph_format.line_spacing = 1.5
    p.alignment = 1

    run = p.add_run('Captions')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = docx.shared.Pt(14)

    p = episode_doc.add_paragraph()
    p.paragraph_format.space_after = 0
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = docx.shared.Mm(12.5)
    p.alignment = 3

    run = p.add_run(transcription_text)
    run.font.name = 'Times New Roman'
    run.font.size = docx.shared.Pt(14)

    episode_doc.save(f'{i + 1}. {video_title}.docx')

timestamps_doc.save(f'{len(video_ids)} Timestamps.docx')